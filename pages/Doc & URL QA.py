__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
import streamlit as st
import textwrap
from docx import Document
from langchain_community.document_loaders import TextLoader
from PyPDF2 import PdfReader
import requests
from bs4 import BeautifulSoup
from langchain.text_splitter import CharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEndpoint
import os
from langchain.chains import RetrievalQA

# Helper functions
def wrap_text_preserve_newlines(text, width=110):
    lines = text.split('\n')
    wrapped_lines = [textwrap.fill(line, width=width) for line in lines]
    wrapped_text = '\n'.join(wrapped_lines)
    return wrapped_text

def read_docx(file_path):
    doc = Document(file_path)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

def process_llm_response(llm_response):
    result = llm_response['result']
    return result

# Streamlit interface
st.title("Files and URL Q/A")

# Option to choose between file upload or URL input
option = st.radio("Choose an option:", ("Upload a file", "Enter a URL"))

document = None

if option == "Upload a file":
    file = st.file_uploader("Upload a DOCX, PDF, or TXT file", type=["docx", "pdf", "txt"])
    if file:
        if file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            plain_text = read_docx(file)
            temp_txt_path = 'temp_resume.txt'
            with open(temp_txt_path, 'w', encoding='utf-8') as f:
                f.write(plain_text)
            loader = TextLoader(temp_txt_path)
            document = loader.load()
        elif file.type == "application/pdf":
            reader = PdfReader(file)
            pdf_text = ''.join([page.extract_text() or '' for page in reader.pages])
            with open("temp_pdf_text.txt", 'w', encoding='utf-8') as f:
                f.write(pdf_text)
            loader = TextLoader("temp_pdf_text.txt")
            document = loader.load()
        elif file.type == "text/plain":
            txt_text = file.read().decode('utf-8')
            with open("temp_txt_text.txt", 'w', encoding='utf-8') as f:
                f.write(txt_text)
            loader = TextLoader("temp_txt_text.txt")
            document = loader.load()

elif option == "Enter a URL":
    url = st.text_input("Enter a URL")
    if url:
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        web_text = soup.get_text()
        document = [LangchainDocument(page_content=web_text)]

if document:
    # Initialize the text splitter and vector store
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    split_docs = text_splitter.split_documents(document)
    embeddings = HuggingFaceEmbeddings()
    persist_directory = 'db'
    vectordb = Chroma.from_documents(documents = document, embedding = embeddings, persist_directory = persist_directory)
    vectordb.persist()
    vectordb = None
    vectordb = Chroma(persist_directory = persist_directory, embedding_function = embeddings)

    # Query input
    query = st.text_input("Ask questions from your file")
    if query:
        os.environ["HUGGINGFACEHUB_API_TOKEN"] = "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"
        repo_id = "mistralai/Mistral-7B-Instruct-v0.2"
        llm = HuggingFaceEndpoint(repo_id=repo_id, max_length=200, temperature=0.7)
        retriever = vectordb.as_retriever()
        qa_chain = RetrievalQA.from_chain_type(llm=llm,
                                              chain_type="stuff",
                                              retriever=retriever,
                                              return_source_documents=True)
        llm_response = qa_chain(query)
        result = process_llm_response(llm_response)
        st.write("Response:", wrap_text_preserve_newlines(result))
